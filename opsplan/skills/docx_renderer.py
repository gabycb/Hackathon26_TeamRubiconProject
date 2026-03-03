"""DOCX Renderer Skill — Convert SOP JSON to formatted Word document."""
import json
from pathlib import Path
from semantic_kernel.functions import kernel_function


class DocxRendererSkill:

    @kernel_function(name="render_sop_docx", description="Convert a validated SOP JSON document into a formatted Word (.docx) file. Returns the file path.")
    async def render_sop_docx(self, sop_json: str, output_path: str = "mission_plan.docx") -> str:
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            sop = json.loads(sop_json)
            doc = Document()
            doc.add_heading("Team Rubicon — Operations Order", level=0)
            sections = [
                ("I. Situation", sop.get("situation", {})),
                ("II. Mission", sop.get("mission", {})),
                ("III. Execution", sop.get("execution", {})),
                ("IV. Sustainment", sop.get("sustainment", {})),
                ("V. Command & Signal", sop.get("command_signal", {})),
            ]
            for title, content in sections:
                doc.add_heading(title, level=1)
                if isinstance(content, dict):
                    for key, value in content.items():
                        doc.add_heading(key.replace("_", " ").title(), level=2)
                        if isinstance(value, str):
                            doc.add_paragraph(value)
                        elif isinstance(value, list):
                            for item in value:
                                if isinstance(item, dict):
                                    for k, v in item.items():
                                        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
                                else:
                                    doc.add_paragraph(str(item), style="List Bullet")
                        elif isinstance(value, dict):
                            for k, v in value.items():
                                doc.add_paragraph(f"{k}: {v}", style="List Bullet")
            doc.save(output_path)
            return json.dumps({"success": True, "path": output_path})
        except Exception as e:
            return json.dumps({"success": False, "error": str(e)})

    def render(self, sop_dict: dict, output_path: str) -> str:
        """Synchronous render for direct calls."""
        import asyncio
        return asyncio.run(self.render_sop_docx(json.dumps(sop_dict), output_path))
