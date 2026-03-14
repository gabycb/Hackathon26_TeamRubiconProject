"""
Generate a .docx field assessment report.
Uses python-docx (same as export_sop.py).
"""
import io
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def build_assessment_docx(data: dict) -> io.BytesIO:
    if not HAS_DOCX:
        buf = io.BytesIO()
        buf.write(b"python-docx not installed")
        buf.seek(0)
        return buf

    doc = Document()

    # Style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Field Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Pipeline: {data.get('pipeline', 'unknown')}")
    doc.add_paragraph("")

    # Zone info
    zone = data.get("zone", {})
    if zone:
        doc.add_heading("Zone Information", level=1)
        doc.add_paragraph(f"Zone: {zone.get('area_name', zone.get('fips_tract', 'Unknown'))}")
        doc.add_paragraph(f"FIPS Tract: {zone.get('fips_tract', '—')}")
        doc.add_paragraph(f"Composite Score: {zone.get('composite_score', '—')}")
        doc.add_paragraph(f"SVI Score: {zone.get('svi_score', '—')}")
        doc.add_paragraph(f"Population: {zone.get('population', '—')}")

    # Assessment summary
    doc.add_heading("Damage Assessment", level=1)

    classification = data.get("damage_classification", "unknown").upper()
    pct = data.get("damage_percentage", 0)
    p = doc.add_paragraph()
    run = p.add_run(f"Classification: {classification} — {pct}% damage")
    run.bold = True
    run.font.size = Pt(14)
    if classification in ("DESTROYED", "MAJOR"):
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    elif classification == "MINOR":
        run.font.color.rgb = RGBColor(0xCA, 0x8A, 0x04)
    elif classification == "NONE":
        run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    if data.get("summary"):
        doc.add_paragraph(data["summary"])

    if data.get("confidence"):
        doc.add_paragraph(f"AI Confidence: {int(data['confidence'] * 100)}%")

    if data.get("structure_type"):
        doc.add_paragraph(f"Structure Type: {data['structure_type']}")

    if data.get("photos_analyzed"):
        doc.add_paragraph(f"Photos Analyzed: {data['photos_analyzed']} ({data.get('photos_with_structure', '?')} with structure)")

    # Component breakdown table
    components = data.get("components", {})
    if components:
        doc.add_heading("Component Damage Breakdown", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Component"
        hdr[1].text = "Damage Level"
        hdr[2].text = "Notes"
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for comp_name, comp_data in components.items():
            row = table.add_row().cells
            row[0].text = comp_name.replace("_", " ").title()
            if isinstance(comp_data, dict):
                row[1].text = (comp_data.get("damage", "—")).upper()
                row[2].text = comp_data.get("notes", "")
            else:
                row[1].text = str(comp_data)

    # Human review
    rejected = data.get("rejected_components", [])
    if rejected:
        doc.add_heading("Human Review Notes", level=2)
        doc.add_paragraph(f"Field team rejected AI findings for: {', '.join(rejected)}")
    if data.get("human_reviewed"):
        doc.add_paragraph("✓ Assessment was human-reviewed before submission.")

    # Hazards
    hazards = data.get("hazards", [])
    if hazards:
        doc.add_heading("Hazards Identified", level=2)
        for h in hazards:
            doc.add_paragraph(f"⚠ {h}", style="List Bullet")

    # Recommended actions
    actions = data.get("recommended_actions", [])
    if actions:
        doc.add_heading("Recommended Actions", level=2)
        for i, a in enumerate(actions, 1):
            doc.add_paragraph(f"{i}. {a}")

    # Field team tags
    tags = data.get("tags", {})
    if tags:
        doc.add_heading("Field Team Observations", level=2)
        if tags.get("hazards"):
            doc.add_paragraph(f"Hazard tags: {', '.join(tags['hazards'])}")
        if tags.get("damage"):
            doc.add_paragraph(f"Damage tags: {', '.join(tags['damage'])}")
        if tags.get("notes"):
            doc.add_paragraph(f"Notes: {tags['notes']}")

    # Vision metadata
    if data.get("vision_tags"):
        doc.add_heading("AI Vision Metadata", level=2)
        doc.add_paragraph(f"Caption: {data.get('vision_caption', '—')}")
        doc.add_paragraph(f"Tags: {', '.join(data.get('vision_tags', []))}")
        if data.get("ocr_text"):
            doc.add_paragraph(f"OCR Text: {'; '.join(data['ocr_text'])}")

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph("Generated by OpsPlan — THYNK UNLIMITED / Team Rubicon")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
